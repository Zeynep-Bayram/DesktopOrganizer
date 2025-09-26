#!/usr/bin/env python3
"""
İçerik çıkarma modülü - PDF, görsel ve DOCX dosyalarından metin çıkarır
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any

# PDF işleme
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.layout import LAParams
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Uyarı: 'pdfminer.six' paketi yüklü değil. PDF çıkarma devre dışı.")

# DOCX işleme
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Uyarı: 'python-docx' paketi yüklü değil. DOCX çıkarma devre dışı.")

# OCR işleme (pytesseract ve Pillow gerektirir)
try:
    from PIL import Image
    import pytesseract
    # Tesseract OCR motorunun sistem PATH'inde olmaması durumunda aşağıdaki satırı etkinleştirin
    # ve Tesseract'ın yürütülebilir dosyasının yolunu belirtin.
    # Örnek (Windows): pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    # Örnek (Linux/macOS): pytesseract.pytesseract.tesseract_cmd = r'/usr/local/bin/tesseract'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Uyarı: 'pytesseract' veya 'Pillow' paketi yüklü değil. Görselden metin çıkarma devre dışı.")
except Exception as e:
    # Tesseract motorunun kendisi bulunamazsa veya başka bir başlatma hatası olursa
    OCR_AVAILABLE = False
    print(f"Uyarı: Tesseract OCR motoru veya yapılandırmasında hata oluştu: {e}. Görselden metin çıkarma devre dışı.")


logger = logging.getLogger(__name__)

class ContentExtractor:
    """Dosyalardan içerik çıkarma sınıfı"""
    
    def __init__(self):
        self.supported_extensions = {
            'pdf': self._extract_pdf_content,
            'docx': self._extract_docx_content,
            'doc': self._extract_docx_content, 
                                                
            'jpg': self._extract_image_content,
            'jpeg': self._extract_image_content,
            'png': self._extract_image_content,
            'bmp': self._extract_image_content,
            'tiff': self._extract_image_content,
            'gif': self._extract_image_content
        }
        
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        """
        Dosyadan içerik çıkarır
        
        Args:
            file_path: Dosya yolu
            
        Returns:
            Dict içerisinde:
            - success: İşlem başarılı mı
            - content: Çıkarılan metin
            - file_type: Dosya türü
            - error: Hata mesajı (varsa)
        """
        result = {
            'success': False,
            'content': '',
            'file_type': '',
            'error': None
        }
        
        try:
            if not file_path.exists():
                result['error'] = f"Dosya bulunamadı: {file_path}"
                return result
                
            # Dosya uzantısını al
            extension = file_path.suffix.lower().lstrip('.')
            result['file_type'] = extension
            
            # Desteklenen uzantı kontrolü
            if extension not in self.supported_extensions:
                result['error'] = f"Desteklenmeyen dosya türü: {extension}"
                return result
            
            # İlgili çıkarma fonksiyonunu çağır
            extractor_func = self.supported_extensions[extension]
            content = extractor_func(file_path)
            
            if content:
                result['success'] = True
                result['content'] = content
                logger.info(f"İçerik çıkarıldı: {file_path.name} ({len(content)} karakter)")
            else:
                result['error'] = "İçerik çıkarılamadı veya dosya boş."
                
        except Exception as e:
            result['error'] = f"İçerik çıkarma hatası: {str(e)}"
            logger.error(f"İçerik çıkarma hatası ({file_path.name}): {e}")
            
        return result
    
    def _extract_pdf_content(self, file_path: Path) -> Optional[str]:
        """PDF dosyasından metin çıkarır"""
        if not PDF_AVAILABLE:
            logger.error("pdfminer.six paketi yüklü değil veya başlatılamadı.")
            return None
            
        try:
            # Dosyayı copy edip temp'den oku (kilitleme önlemi)
            import tempfile
            import shutil
            
            # Orijinal dosyanın uzantısını kullanarak tempfile oluştur
            suffix = file_path.suffix
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                shutil.copy2(str(file_path), temp_file.name)
                temp_path = temp_file.name
            
            try:
                # Temp dosyayı oku
                with open(temp_path, 'rb') as file:
                    # LAParams ile daha iyi metin çıkarma
                    laparams = LAParams(
                        all_texts=True,
                        word_margin=0.1,
                        char_margin=2.0,
                        line_margin=0.5,
                        boxes_flow=0.5
                    )
                    
                    text = pdf_extract_text(file, laparams=laparams)
                
                # Metni temizle
                if text:
                    text = text.strip()
                    # Fazla boşlukları tek boşluğa indirge
                    text = ' '.join(text.split())
                    
                    return text
                    
            finally:
                # Temp dosyayı sil
                try:
                    # Dosya handle'ının tamamen bırakıldığından emin olmak için küçük bir gecikme
                    import time
                    time.sleep(0.05)
                    Path(temp_path).unlink()
                except OSError as os_err:
                    logger.warning(f"Geçici PDF dosyası silinirken hata oluştu ({temp_path}): {os_err}")
                
        except Exception as e:
            logger.error(f"PDF çıkarma hatası ({file_path.name}): {e}")
            
        return None
    
    def _extract_docx_content(self, file_path: Path) -> Optional[str]:
        """DOCX dosyasından metin çıkarır"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx paketi yüklü değil veya başlatılamadı.")
            return None
            
        try:
            # Dosyayı copy edip temp'den oku (kilitleme önlemi)
            import tempfile
            import shutil
            
            suffix = file_path.suffix
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                shutil.copy2(str(file_path), temp_file.name)
                temp_path = temp_file.name
            
            try:
                # Temp dosyayı oku
                doc = Document(temp_path)
                
                # Tüm paragrafları birleştir
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                # Tabloları da ekle
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            paragraphs.append(' | '.join(row_text))
                
                result = '\n'.join(paragraphs) if paragraphs else None
                
                return result
                
            finally:
                # Temp dosyayı sil
                try:
                    import time
                    time.sleep(0.05)
                    Path(temp_path).unlink()
                except OSError as os_err:
                    logger.warning(f"Geçici DOCX dosyası silinirken hata oluştu ({temp_path}): {os_err}")
            
        except Exception as e:
            logger.error(f"DOCX çıkarma hatası ({file_path.name}): {e}")
            
        return None
    
    def _extract_image_content(self, file_path: Path) -> Optional[str]:
        """Görsel dosyasından OCR ile metin çıkarır"""
        if not OCR_AVAILABLE:
            logger.error("pytesseract veya Pillow paketi yüklü değil veya Tesseract motoru bulunamadı.")
            return None
            
        try:
            # Resmi aç
            image = Image.open(str(file_path))
            
            # OCR ile metin çıkar (Türkçe + İngilizce)
            # config='--psm 6' genellikle iyi çalışır, metin blokları için
            # Daha fazla bilgi: https://tesseract-ocr.github.io/tessdoc/Command-Line-Usage.html#page-segmentation-modes
            text = pytesseract.image_to_string(
                image, 
                lang='tur+eng',  # Türkçe ve İngilizce
                config='--psm 6' # Uniform text block (tek bir sütun metin gibi)
            )
            
            # Metni temizle
            if text:
                text = text.strip()
                # Fazla boşlukları ve satır sonlarını tek boşluğa indirge
                text = ' '.join(text.split())
                return text if len(text) > 3 else None  # En az 3 karakterden kısa metinleri yok say
            return None
                
        except Exception as e:
            logger.error(f"OCR çıkarma hatası ({file_path.name}): {e}")
            
        return None
    
    def is_supported(self, file_path: Path) -> bool:
        """Dosya türünün desteklenip desteklenmediğini kontrol eder"""
        extension = file_path.suffix.lower().lstrip('.')
        return extension in self.supported_extensions
    
    def get_supported_extensions(self) -> list:
        """Desteklenen dosya uzantılarını döndürür"""
        return list(self.supported_extensions.keys())